#!/usr/bin/env python3
"""
build_report.py — generate the Founder OS final-year project report as a .docx.

It reads the editable prose in ``content/*.md`` (a small Markdown subset),
embeds rendered diagrams from ``figures/*.png``, and applies the college
formatting spec exactly:

    * A4 paper, Times New Roman 12 pt, 1.5 line spacing, fully justified
    * 1.5" left margin (binding) + 1" top/bottom/right
    * lowercase-roman page numbers on the preliminary pages, then Arabic
      numerals restarting at 1 from Chapter 1
    * auto Table of Contents + List of Figures + List of Tables (Word fields)

Usage:
    python report/build_report.py
Output:
    report/Founder_OS_Report.docx

Edit the CONFIG block below with your real title-page / certificate details.
Anything left as «...» is a placeholder you should replace.
"""

import os
import re
import glob

from docx import Document
from docx.shared import Pt, Inches, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
CONTENT = os.path.join(HERE, "content")
FIGURES = os.path.join(HERE, "figures")
OUTPUT = os.path.join(HERE, "Founder_OS_Report.docx")

# ───────────────────────── CONFIG: title page + certificate ─────────────────
# Replace every «...» with your real details. Lines you leave as «...» will
# show up verbatim so you can spot what still needs filling.
CONFIG = {
    "title": "FOUNDER OS",
    "subtitle": "An AI-Augmented Operations Intelligence Platform "
                "for Multi-Startup Portfolio Management",
    "degree": "«Bachelor of Technology»",
    "branch": "«Computer Science and Engineering»",
    "student_name": "«STUDENT NAME»",
    "roll_no": "«ROLL / REGISTRATION NO.»",
    "guide_name": "«GUIDE NAME»",
    "guide_designation": "«Assistant Professor, Dept. of CSE»",
    "hod_name": "«HEAD OF DEPARTMENT NAME»",
    "department": "«Department of Computer Science and Engineering»",
    "college": "«COLLEGE / INSTITUTE NAME»",
    "university": "«UNIVERSITY NAME»",
    "place": "«CITY»",
    "academic_year": "«2025–2026»",
    "month_year": "«JUNE 2026»",
    "logo_path": os.path.join(FIGURES, "logo.png"),  # optional; skipped if absent
}

# Order of preliminary (roman-numbered) markdown files.
PRELIM_FILES = [
    "02_acknowledgements.md",
    "03_abstract.md",
    "04_nomenclature.md",
]
# Order of body (decimal-numbered) markdown files.
BODY_FILES = [
    "10_intro.md",
    "20_litreview.md",
    "30_design.md",
    "40_implementation.md",
    "50_results.md",
    "60_conclusion.md",
    "70_references.md",
    "80_appendix.md",
]

TNR = "Times New Roman"
MONO = "Consolas"


# ───────────────────────────── low-level helpers ────────────────────────────
def _el(tag):
    return OxmlElement(tag)


def add_field(paragraph, instr, default_text=""):
    """Insert a Word field (e.g. PAGE, TOC, SEQ) into a paragraph."""
    run = paragraph.add_run()
    begin = _el("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr_el = _el("w:instrText"); instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    sep = _el("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = _el("w:t"); t.text = default_text
    end = _el("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr_el, sep, t, end):
        run._r.append(node)
    return run


def shade(element_pr, hex_fill):
    """Attach a <w:shd> fill to a pPr or tcPr element."""
    shd = _el("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    element_pr.append(shd)


def shade_cell(cell, hex_fill):
    shade(cell._tc.get_or_add_tcPr(), hex_fill)


def shade_paragraph(paragraph, hex_fill):
    shade(paragraph._p.get_or_add_pPr(), hex_fill)


def set_pgnum(section, fmt, start=None):
    """Set the page-number numeral format (lowerRoman | decimal) for a section."""
    sectPr = section._sectPr
    for existing in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(existing)
    pg = _el("w:pgNumType")
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))
    sectPr.append(pg)


def footer_pagenum(section):
    """Centered 'Page N' field in the section footer (unlinked from previous)."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Page ")
    run.font.name = TNR
    run.font.size = Pt(10)
    add_field(p, "PAGE", "1")
    for r in p.runs:
        r.font.name = TNR
        r.font.size = Pt(10)


def enable_update_fields(doc):
    """Ask Word to refresh all fields (TOC/LOF/LOT/page refs) when opened."""
    settings = doc.settings.element
    upd = _el("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


# ───────────────────────────── styles ───────────────────────────────────────
def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = TNR
    normal.font.size = Pt(12)
    rpr = normal.element.get_or_add_rPr().get_or_add_rFonts()
    rpr.set(qn("w:ascii"), TNR)
    rpr.set(qn("w:hAnsi"), TNR)
    rpr.set(qn("w:cs"), TNR)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(6)

    sizes = {"Heading 1": 16, "Heading 2": 14, "Heading 3": 12.5}
    for name, size in sizes.items():
        st = doc.styles[name]
        st.font.name = TNR
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        st.paragraph_format.keep_with_next = True


def configure_section(section, prelim=False):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)


# ───────────────────────────── inline markdown ──────────────────────────────
INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*.+?\*)")


def add_inline(paragraph, text):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = MONO; r.font.size = Pt(10.5)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = paragraph.add_run(part[1:-1]); r.italic = True
        else:
            paragraph.add_run(part)


# ───────────────────────────── block builders ───────────────────────────────
def add_caption(doc, label, text, center=True):
    p = doc.add_paragraph()
    try:
        p.style = doc.styles["Caption"]
    except KeyError:
        pass
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(f"{label} ")
    run.bold = True
    add_field(p, f"SEQ {label} \\* ARABIC", "1")
    rest = p.add_run(f": {text}")
    rest.italic = True
    for r in p.runs:
        r.font.name = TNR
        r.font.size = Pt(11)


def render_figure(doc, rel_path, caption):
    path = rel_path if os.path.isabs(rel_path) else os.path.join(HERE, rel_path)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(5.6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # placeholder box so the build never fails on a missing render
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        shade_cell(cell, "FDF2D0")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mmd = os.path.basename(rel_path).replace(".png", ".mmd")
        r = para.add_run(f"[ Figure not yet rendered — run render_figures.ps1 to "
                         f"produce {os.path.basename(rel_path)} from figures/{mmd} ]")
        r.italic = True
        r.font.size = Pt(11)
    add_caption(doc, "Figure", caption)


def render_table(doc, rows, caption=None):
    if caption:
        add_caption(doc, "Table", caption, center=False)
    ncol = max(len(r) for r in rows)
    tbl = doc.add_table(rows=0, cols=ncol)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        cells = tbl.add_row().cells
        for j in range(ncol):
            val = row[j] if j < len(row) else ""
            para = cells[j].paragraphs[0]
            para.text = ""
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.space_after = Pt(2)
            add_inline(para, val)
            for r in para.runs:
                r.font.size = Pt(10.5)
                if i == 0:
                    r.bold = True
            if i == 0:
                shade_cell(cells[j], "D9D9D9")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def render_code(doc, lines):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.left_indent = Inches(0.15)
    pf.line_spacing = 1.0
    pf.space_before = Pt(4)
    pf.space_after = Pt(8)
    shade_paragraph(p, "F4F4F4")
    for i, ln in enumerate(lines):
        r = p.add_run(ln.replace("\t", "    "))
        r.font.name = MONO
        r.font.size = Pt(9.5)
        if i < len(lines) - 1:
            r.add_break()


def add_toc_field(doc, heading, instr):
    h = doc.add_paragraph(heading, style="Heading 1")
    h.paragraph_format.page_break_before = False
    p = doc.add_paragraph()
    add_field(p, instr, "Right-click → Update Field to populate.")


# ───────────────────────────── markdown renderer ────────────────────────────
def render_markdown(doc, text, page_break_before_h1=True):
    lines = text.split("\n")
    i = 0
    seen_h1 = False
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # fenced code block
        if stripped.startswith("```"):
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            render_code(doc, code)
            i += 1
            continue

        # directives
        if stripped == "[[pagebreak]]":
            doc.add_page_break()
            i += 1
            continue

        # figure: ![caption](path)
        m = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if m:
            render_figure(doc, m.group(2), m.group(1))
            i += 1
            continue

        # table caption directive: "Table: ..."
        caption = None
        if stripped.startswith("Table:"):
            caption = stripped[len("Table:"):].strip()
            i += 1
            # skip blank lines
            while i < len(lines) and not lines[i].strip():
                i += 1
            stripped = lines[i].strip() if i < len(lines) else ""

        # table block — a run of consecutive lines starting with "|".
        # The first non-separator row is treated as the header; a markdown
        # separator row (e.g. |---|---|) is optional and skipped if present.
        if stripped.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rowline = lines[i].strip()
                if re.match(r"^\|[\s:\-|]+\|?$", rowline):  # separator row
                    i += 1
                    continue
                rows.append([c.strip() for c in rowline.strip("|").split("|")])
                i += 1
            if rows:
                render_table(doc, rows, caption)
            continue

        # headings
        if stripped.startswith("### "):
            doc.add_paragraph(stripped[4:], style="Heading 3")
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_paragraph(stripped[3:], style="Heading 2")
            i += 1
            continue
        if stripped.startswith("# "):
            h = doc.add_paragraph(stripped[2:], style="Heading 1")
            if page_break_before_h1 and seen_h1:
                h.paragraph_format.page_break_before = True
            seen_h1 = True
            i += 1
            continue

        # bullet / numbered lists
        if re.match(r"^[-*] ", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
            i += 1
            continue
        if re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s*", "", stripped))
            i += 1
            continue

        # blockquote → indented note
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run(stripped[2:])
            r.italic = True
            i += 1
            continue

        # horizontal rule → skip
        if stripped == "---":
            i += 1
            continue

        # blank line
        if not stripped:
            i += 1
            continue

        # paragraph
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1


# ───────────────────────────── preliminary pages ────────────────────────────
def centered(doc, text, size=12, bold=False, italic=False, space=6, caps=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text.upper() if caps else text)
    r.bold = bold
    r.italic = italic
    r.font.name = TNR
    r.font.size = Pt(size)
    return p


def build_title_page(doc, cfg):
    centered(doc, cfg["title"], size=30, bold=True, space=4)
    centered(doc, cfg["subtitle"], size=14, italic=True, space=14)
    centered(doc, "A Project Report submitted in partial fulfilment of the "
                  "requirements for the award of the degree of", size=12, space=10)
    centered(doc, f"{cfg['degree']}", size=15, bold=True, space=2)
    centered(doc, f"in {cfg['branch']}", size=13, space=14)
    centered(doc, "Submitted by", size=12, space=4)
    centered(doc, cfg["student_name"], size=15, bold=True, space=2)
    centered(doc, f"Roll No.: {cfg['roll_no']}", size=12, space=14)
    centered(doc, "Under the guidance of", size=12, space=4)
    centered(doc, cfg["guide_name"], size=14, bold=True, space=2)
    centered(doc, cfg["guide_designation"], size=12, space=14)

    if os.path.exists(cfg["logo_path"]):
        doc.add_picture(cfg["logo_path"], width=Inches(1.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    centered(doc, cfg["department"], size=13, bold=True, space=2)
    centered(doc, cfg["college"], size=13, bold=True, space=2)
    centered(doc, cfg["university"], size=12, space=2)
    centered(doc, cfg["month_year"], size=12, bold=True, space=2)


def build_certificate(doc, cfg):
    doc.add_page_break()
    centered(doc, "BONAFIDE CERTIFICATE", size=16, bold=True, space=18)
    body = (
        f"This is to certify that the project report entitled "
        f"“{cfg['title']}: {cfg['subtitle']}” is a bonafide record of the "
        f"work carried out by {cfg['student_name']} (Roll No. {cfg['roll_no']}) in "
        f"partial fulfilment of the requirements for the award of the degree of "
        f"{cfg['degree']} in {cfg['branch']} of {cfg['university']}, during the "
        f"academic year {cfg['academic_year']}. The work presented in this report "
        f"is original and has not been submitted, in part or in full, to any other "
        f"institute or university for the award of any degree or diploma."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    add_inline(p, body)

    doc.add_paragraph().paragraph_format.space_after = Pt(48)
    doc.add_paragraph().paragraph_format.space_after = Pt(48)

    tbl = doc.add_table(rows=2, cols=2)
    tbl.autofit = True
    data = [
        (f"{cfg['guide_name']}", f"{cfg['hod_name']}"),
        (f"Project Guide\n{cfg['guide_designation']}", f"Head of Department\n{cfg['department']}"),
    ]
    for ri, row in enumerate(data):
        for ci, val in enumerate(row):
            cell = tbl.cell(ri, ci)
            cell.paragraphs[0].text = ""
            for k, ln in enumerate(val.split("\n")):
                para = cell.paragraphs[0] if k == 0 else cell.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                rr = para.add_run(ln)
                rr.font.name = TNR
                rr.font.size = Pt(12)
                if ri == 0:
                    rr.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    p2 = doc.add_paragraph()
    p2.add_run("Submitted for the project viva-voce examination held on ____________________.")
    doc.add_paragraph().paragraph_format.space_after = Pt(36)
    e = doc.add_paragraph()
    e.alignment = WD_ALIGN_PARAGRAPH.LEFT
    e.add_run("Internal Examiner                                        External Examiner")


def read_md(name):
    path = os.path.join(CONTENT, name)
    with open(path, encoding="utf-8") as f:
        return f.read()


# ───────────────────────────── assembly ─────────────────────────────────────
def main():
    doc = Document()
    setup_styles(doc)

    # Section 1 — preliminary pages (roman numerals)
    sec1 = doc.sections[0]
    configure_section(sec1, prelim=True)
    sec1.different_first_page_header_footer = True  # no number on title page
    set_pgnum(sec1, "lowerRoman", start=1)
    footer_pagenum(sec1)

    build_title_page(doc, CONFIG)
    build_certificate(doc, CONFIG)

    for fname in PRELIM_FILES:
        doc.add_page_break()
        render_markdown(doc, read_md(fname), page_break_before_h1=False)

    # Table of Contents + List of Figures + List of Tables
    doc.add_page_break()
    add_toc_field(doc, "TABLE OF CONTENTS", 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_page_break()
    add_toc_field(doc, "LIST OF FIGURES", 'TOC \\h \\z \\c "Figure"')
    doc.add_page_break()
    add_toc_field(doc, "LIST OF TABLES", 'TOC \\h \\z \\c "Table"')

    # Section 2 — body (decimal numerals, restart at 1)
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body_section, prelim=False)
    body_section.different_first_page_header_footer = False
    set_pgnum(body_section, "decimal", start=1)
    footer_pagenum(body_section)

    for idx, fname in enumerate(BODY_FILES):
        render_markdown(doc, read_md(fname), page_break_before_h1=(idx > 0))

    enable_update_fields(doc)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
